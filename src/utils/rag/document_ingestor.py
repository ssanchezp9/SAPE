import os
from pathlib import Path
from dotenv import load_dotenv
from typing import Optional, Dict, Union, List
from logging import Logger

from docx import Document as DocxDocument
from langchain_core.documents import Document as LCDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from langchain_postgres.vectorstores import PGVector

from src.utils.database.database import Database
from src.utils.database.dsn_loader import load_dsn
from src.utils.logger import logging_config
from src.utils.env.env import EnvConfig

# Variable global para el modelo de embeddings
EMBEDDING_MODEL = "text-embedding-3-large"
COLLECTION_NAME = "documentos"
DIMENSION = 3072

class DocumentIngestor():
    def __init__(
        self,
    ):        
        self.openai_key: str = EnvConfig().get_open_api_key()
        self.logger: Logger = logging_config.get_logger(__name__)

        self.embeddings = OpenAIEmbeddings(
            model=EMBEDDING_MODEL,
            openai_api_key=self.openai_key
        )

        connection_string = load_dsn()
        self.logger.info(f"Conexión a la base de datos: {connection_string}")
        

        self.vectorstore = PGVector(
            embeddings=self.embeddings,
            collection_name=COLLECTION_NAME,
            connection=connection_string,
            use_jsonb=True,
            create_extension=False,
            collection_metadata={"dimensions": DIMENSION}
        )

    def ingest(self, file_path: str, metadata: Optional[Dict] = None):
        ext = Path(file_path).suffix.lower()
        self.logger.info(f"📥 Iniciando ingesta del archivo: {file_path}")
        if ext == ".docx":
            text = self._load_docx(file_path)
        else:
            raise ValueError(f"Tipo de archivo no soportado: {ext}")

        documents = self._split_text(text, metadata or {"source": Path(file_path).name})
        self.vectorstore.add_documents(documents)
        self.logger.info(f"✅ Ingesta completada: {file_path}")

    def _load_docx(self, file_path: str) -> str:
        doc = DocxDocument(file_path)
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    def _split_text(self, text: str, metadata: Dict) -> list:
        splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
        chunks = splitter.split_text(text)
        return [LCDocument(page_content=chunk, metadata=metadata) for chunk in chunks]
